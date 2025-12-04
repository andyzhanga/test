import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import datetime
import os
import matplotlib as mpl

# 尝试设置全局字体，确保支持中文。
# 你需要安装一个支持中文的字体，例如 'SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei' 等。
# 查找你系统中的字体文件路径，或者直接使用字体名称。
try:
    mpl.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei', 'Arial Unicode MS'] # 尝试多个字体
    mpl.rcParams['axes.unicode_minus'] = False # 解决负号显示问题
except Exception as e:
    print(f"Warning: Font setting failed, may cause display issues: {e}")

# ... 你的绘图代码 ...
plt.title("中文标题")
plt.xlabel("X轴标签")
plt.ylabel("Y轴标签")
print("Matplotlib configuration for Chinese characters applied.")

# 在保存图片时，也可以指定字体的编码，但通常全局设置更方便
# plt.savefig('plot.png', dpi=300, bbox_inches='tight', font_properties=font_prop)


# --- 配置区 ---
EXCEL_FILE_PATH = 'security_log.xls'  # 您的日志 .xls 文件路径
REPORT_OUTPUT_PATH = 'security_analysis_report.docx' # 生成的Word报告路径

# 确保你的 Excel 文件在脚本运行的目录下，或者提供完整路径
# 如果文件不存在，此脚本会报错退出。

# --- 日志列名映射 (请根据你的xls文件的实际列名修改) ---
# Key 是Excel中的原始列名, Value 是我们在脚本中使用的标准化列名
COLUMN_MAPPING = {
    '日期/时间': '时间',             # 假设Excel中有这个列，如果只是时间，可以只写 '时间'
    '时间': '时间',                 # 有些文件可能只有时间列
    '威胁类型': '威胁类型',
    '严重性': '严重性',
    '威胁名称': '威胁名称',
    '次数': '次数',                  # 如果是单条日志Denne，这个可以忽略或设为1；如果Excel本身就有次数统计，则保留；否则可以删除
    '源安全区域': '源安全区域',
    '目的安全区域': '目的安全区域',
    '攻击者': '攻击者',               # 攻击源 IP 或主机名
    '攻击目标': '攻击目标',           # 攻击目标 IP 或主机名
    '源端口': '源端口',
    '源地区': '源地区',
    '目的端口': '目的端口',
    '目的地区': '目的地区',
    '应用': '应用',
    '协议': '协议',
    '动作': '动作',
}

# --- 辅助函数 ---

def create_plot_and_save(data, x_col, y_col, title, plot_filename, top_n=10):
    """
    生成柱状图并保存，返回图片路径
    """
    plt.figure(figsize=(10, 6))
    plot_data = data.nlargest(top_n, y_col) if y_col else data.head(top_n)
    sns.barplot(x=y_col, y=x_col, data=plot_data)
    plt.title(title)
    plt.tight_layout()
    plot_path = plot_filename
    plt.savefig(plot_path)
    plt.close()
    return plot_path

def create_pie_chart_and_save(data, value_col, title, plot_filename):
    """
    生成饼状图并保存，返回图片路径
    """
    plt.figure(figsize=(7, 7))
    # 确保value_col是DataFrame的列名
    if isinstance(data, pd.Series):
        counts = data.value_counts()
    elif isinstance(data, pd.DataFrame):
        counts = data[value_col].value_counts()
    else:
        raise TypeError("Data must be a pandas Series or DataFrame")

    counts.plot.pie(autopct='%1.1f%%', startangle=90)
    plt.title(title)
    plt.ylabel('') # Hide the y-axis label for pie chart
    plt.tight_layout()
    plot_path = plot_filename
    plt.savefig(plot_path)
    plt.close()
    return plot_path

def add_table_to_document(doc, dataframe, caption):
    """
    将DataFrame添加到Word文档中作为表格
    """
    doc.add_paragraph(caption)
    # 确保DataFrame不为空
    if dataframe.empty:
        doc.add_paragraph("  (无数据)")
        return

    table = doc.add_table(rows=dataframe.shape[0] + 1, cols=dataframe.shape[1])
    table.style = 'Table Grid' # 使用清晰的表格样式

    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(dataframe.columns):
        hdr_cells[i].text = col_name

    # 添加数据行
    for i in range(dataframe.shape[0]):
        row_cells = table.rows[i + 1].cells
        # 确保所有数据都是字符串
        for j, col_name in enumerate(dataframe.columns):
            row_cells[j].text = str(dataframe.iloc[i, j])

# --- 主处理函数 ---

def generate_security_report(excel_file, output_docx_file):
    """
    读取Excel日志，进行分析，生成Word报告
    """
    if not os.path.exists(excel_file):
        print(f"错误：Excel 文件 '{excel_file}' 不存在。请确保文件路径正确。")
        return

    try:
        # 1. 读取Excel文件
        print(f"正在读取 Excel 文件: '{excel_file}'...")
        # 注意：如果你读取 .xls 文件遇到问题，可能需要安装 xlrd: pip install xlrd
        # pandas 默认会尝试使用 'openpyxl' 来读取 xlsx, 可以显式指定 engine='xlrd' 来读取xls
        df = pd.read_excel(excel_file, engine='xlrd') # explicitly use xlrd for .xls files
        print("Excel 文件读取成功。")

        # 2. 列名标准化
        # 检查并替换列名，使其与COLUMN_MAPPING中的Keys匹配
        # 找出Excel中存在的列，并映射到标准化列名
        actual_columns = {}
        for excel_col, std_col in COLUMN_MAPPING.items():
            if excel_col in df.columns:
                actual_columns[excel_col] = std_col
                # 将Excel中的列名重命名为标准列名
                df.rename(columns={excel_col: std_col}, inplace=True)

        # 检查是否所有必要的列都存在
        required_cols = ['时间', '威胁类型', '严重性', '威胁名称', '源地区', '攻击目标', '动作']
        missing_cols = [col for col in required_cols if col not in df.columns]
        if missing_cols:
            print(f"警告：Excel文件中缺少以下关键列：{', '.join(missing_cols)}")
            print("脚本可能无法正常工作。请检查 COLUMN_MAPPING 配置。")
            # 根据需要可以选择在此处退出或继续（部分功能可能受限）
            # return

        # 3. 数据预处理与清洗
        print("正在进行数据预处理...")
        # 转换时间列为 datetime 对象
        if '时间' in df.columns:
            # 尝试保留原始时间格式，如果它已经是字符串
            if pd.api.types.is_datetime64_any_dtype(df['时间']):
                pass # 已经是datetime类型
            else:
                # 尝试解析多种可能的日期时间格式
                try:
                    df['时间'] = pd.to_datetime(df['时间'], errors='coerce')
                except Exception as e:
                    print(f"警告：无法将‘时间’列统一解析为日期时间类型，请检查格式：{e}")
                    # 如果解析失败，可能需要根据具体格式调整

            # 移除时间列解析失败（NaT）的行
            df.dropna(subset=['时间'], inplace=True)
        else:
            print("警告：'时间'列不存在，无法进行时间相关的排序和统计。")


        # 确保“次数”列是数字类型，如果不存在则创建
        if '次数' not in df.columns:
            df['次数'] = 1 # 默认每行代表一次事件

        # 确保“次数”列是整数，并处理可能的NaN
        df['次数'] = pd.to_numeric(df['次数'], errors='coerce').fillna(1).astype(int)

        # 过滤掉一些可能无意义的日志（例如，没有明确威胁类型的）
        df = df[df['威胁类型'].notna() & (df['威胁类型'] != '')]
        df = df[df['威胁类型'] != '未知'] # 假设‘未知’不是我们关心的具体威胁

        # 排序以便分析时间趋势和展示示例
        df.sort_values(by='时间', inplace=True)

        # 4. 生成Word文档
        document = Document()
        document.add_heading('安全威胁日志分析报告', level=0)

        # --- 报告元信息 ---
        start_time_str = df['时间'].min().strftime('%Y-%m-%d %H:%M:%S') if not df.empty else 'N/A'
        end_time_str = df['时间'].max().strftime('%Y-%m-%d %H:%M:%S') if not df.empty else 'N/A'
        report_gen_time_str = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        total_logs = len(df)

        document.add_paragraph(f"**报告周期：** {start_time_str} 至 {end_time_str}")
        document.add_paragraph(f"**报告生成日期：** {report_gen_time_str}")
        document.add_paragraph(f"**原始日志总数：** {total_logs}")
        document.add_paragraph("---") # 分隔线

        # --- 2. 日志总览 ---
        document.add_heading('2. 日志总览', level=1)
        document.add_paragraph(f"本次分析覆盖从 **{start_time_str}** 到 **{end_time_str}** 的日志。")
        document.add_paragraph(f"共处理了 **{total_logs}** 条有效安全威胁日志。")
        document.add_paragraph("---")

        # --- 3. 威胁类型分析 ---
        document.add_heading('3. 威胁类型分析', level=1)
        if '威胁类型' in df.columns and not df['威胁类型'].empty:
            threat_type_counts = df['威胁类型'].value_counts().reset_index()
            threat_type_counts.columns = ['威胁类型', '次数']

            # 计算占比
            total_threats = threat_type_counts['次数'].sum()
            threat_type_counts['占比'] = (threat_type_counts['次数'] / total_threats * 100).round(2).astype(str) + '%'

            plot_path = create_plot_and_save(threat_type_counts, '威胁类型', '次数', '威胁类型分布 (Top 10)', 'threat_type_distribution.png')
            document.add_paragraph("3.1 威胁类型分布：")
            document.add_picture(plot_path, width=Inches(6.0))
            os.remove(plot_path) # 清理图片文件

            # 添加威胁类型表格 (Top 10)
            add_table_to_document(document, threat_type_counts.head(10), "3.2 常见威胁类型详情 (Top 10):")
        else:
            document.add_paragraph("3.1 威胁类型分布： (无法解析，可能未提供 '威胁类型' 列)")
        document.add_paragraph("---")

        # --- 4. 严重性分析 ---
        document.add_heading('4. 严重性分析', level=1)
        if '严重性' in df.columns and not df['严重性'].empty:
            severity_counts = df['严重性'].value_counts().reset_index()
            severity_counts.columns = ['严重性', '次数']

            plot_path = create_pie_chart_and_save(severity_counts, '严重性', '威胁严重性分布', 'severity_distribution.png')
            document.add_picture(plot_path, width=Inches(4.0))
            os.remove(plot_path)

            # 添加严重性表格
            add_table_to_document(document, severity_counts, "4.1 威胁严重性详情:")
        else:
            document.add_paragraph("4.1 威胁严重性分布： (无法解析，可能未提供 '严重性' 列)")
        document.add_paragraph("---")

        # --- 5. 攻击源分析 ---
        document.add_heading('5. 攻击源分析', level=1)
        if '源地区' in df.columns and not df['源地区'].empty:
            source_region_counts = df['源地区'].value_counts().reset_index()
            source_region_counts.columns = ['源地区', '次数']

            plot_path = create_plot_and_save(source_region_counts, '源地区', '次数', '攻击源地区分布 (Top 10)', 'source_region_distribution.png')
            document.add_paragraph("5.1 攻击源地区分布：")
            document.add_picture(plot_path, width=Inches(6.0))
            os.remove(plot_path)
        else:
            document.add_paragraph("5.1 攻击源地区分布： (无法解析，可能未提供 '源地区' 列)")

        if '攻击者' in df.columns and not df['攻击者'].empty:
            attackers_counts = df['攻击者'].value_counts().reset_index()
            attackers_counts.columns = ['攻击者IP', '次数']
            add_table_to_document(document, attackers_counts.head(10), "5.2 活跃攻击源IP (Top 10):")
        else:
            document.add_paragraph("5.2 活跃攻击源IP： (无法解析，可能未提供 '攻击者' 列)")
        document.add_paragraph("---")

        # --- 6. 攻击目标分析 ---
        document.add_heading('6. 攻击目标分析', level=1)
        if '攻击目标' in df.columns and not df['攻击目标'].empty:
            target_counts = df['攻击目标'].value_counts().reset_index()
            target_counts.columns = ['攻击目标', '次数']
            add_table_to_document(document, target_counts.head(10), "6.1 目标Top 10 (被攻击次数):")
        else:
            document.add_paragraph("6.1 目标Top 10： (无法解析，可能未提供 '攻击目标' 列)")
        document.add_paragraph("---")

        # --- 7. 应用与协议分析 ---
        document.add_heading('7. 应用与协议分析', level=1)
        if '应用' in df.columns and not df['应用'].empty:
            app_counts = df['应用'].value_counts().reset_index()
            app_counts.columns = ['应用', '次数']
            add_table_to_document(document, app_counts.head(10), "7.1 常见应用 (Top 10):")
        else:
            document.add_paragraph("7.1 常见应用： (无法解析，可能未提供 '应用' 列)")

        if '协议' in df.columns and not df['协议'].empty:
            protocol_counts = df['协议'].value_counts().reset_index()
            protocol_counts.columns = ['协议', '次数']
            add_table_to_document(document, protocol_counts.head(10), "7.2 常见协议 (Top 10):")
        else:
            document.add_paragraph("7.2 常见协议： (无法解析，可能未提供 '协议' 列)")
        document.add_paragraph("---")

        # --- 8. 动作分析 ---
        document.add_heading('8. 动作分析', level=1)
        if '动作' in df.columns and not df['动作'].empty:
            action_counts = df['动作'].value_counts().reset_index()
            action_counts.columns = ['动作', '次数']

            plot_path = create_pie_chart_and_save(action_counts, '动作', '安全动作分布', 'action_distribution.png')
            document.add_picture(plot_path, width=Inches(4.0))
            os.remove(plot_path)

            add_table_to_document(document, action_counts, "8.1 动作详情:")
        else:
            document.add_paragraph("8.1 动作分布： (无法解析，可能未提供 '动作' 列)")
        document.add_paragraph("---")

        # --- 9. 威胁家族/名称分析 ---
        document.add_heading('9. 威胁家族/名称分析', level=1)
        if '威胁名称' in df.columns and not df['威胁名称'].empty:
            threat_name_counts = df['威胁名称'].value_counts().reset_index()
            threat_name_counts.columns = ['威胁名称', '次数']
            add_table_to_document(document, threat_name_counts.head(10), "9.1 最常见的威胁名称 (Top 10):")
        else:
            document.add_paragraph("9.1 最常见的威胁名称： (无法解析，可能未提供 '威胁名称' 列)")
        document.add_paragraph("---")

        # --- 10. 重点攻击事件示例 (例如，按严重性排序，取前3个高危事件) ---
        document.add_heading('10. 重点攻击事件示例', level=1)
        if '严重性' in df.columns:
            # 尝试根据严重性排序，将'高'放在前面，然后按时间排序
            severity_order = ['高', '中', '低'] # 假设严重性只有这三档
            if all(s in df['严重性'].unique() for s in severity_order):
                # 创建一个映射，处理可能不存在的严重性级别，避免出错
                severity_map = {level: i for i, level in enumerate(severity_order)}
                # 尝试为所有唯一值提供一个默认值（例如，排在最后）
                max_priority = len(severity_order)
                df['severity_priority'] = df['严重性'].map(severity_map).fillna(max_priority)
                
                # 按优先级和时间排序
                sorted_df = df.sort_values(by=['severity_priority', '时间'])
            else:
                print("警告：严重性级别不符合预期，将按时间排序。")
                sorted_df = df.sort_values(by='时间')
        else:
            sorted_df = df.sort_values(by='时间')

        sample_events = sorted_df.head(5).to_dict('records') # 取前5个
        if sample_events:
            for i, event in enumerate(sample_events):
                document.add_heading(f'10.{i+1} 示例事件', level=2)
                document.add_paragraph(f"  - **时间:** {event.get('时间', 'N/A')}")

                # 尝试获取所有可能要显示的字段
                fields_to_show = ['威胁类型', '严重性', '威胁名称', '次数', '源安全区域', '目的安全区域',
                                  '攻击者', '攻击目标', '源端口', '源地区', '目的端口', '目的地区',
                                  '应用', '协议', '动作']
                for field in fields_to_show:
                    if field in event and event[field] not in [None, '', pd.NaT]:
                         # 格式化时间字段，避免显示毫秒
                        if isinstance(event[field], datetime.datetime):
                             document.add_paragraph(f"  - **{field}:** {event[field].strftime('%Y-%m-%d %H:%M:%S')}")
                        else:
                             document.add_paragraph(f"  - **{field}:** {event[field]}")
        else:
            document.add_paragraph("  (无样本事件可展示)")
        document.add_paragraph("---")


        # --- 11. 结论与建议 ---
        document.add_heading('11. 结论与建议', level=1)
        document.add_paragraph(f"11.1 **整体安全态势评估：**")
        document.add_paragraph("    本周期内，观察到 [例如：来自特定地区的扫描活动频繁，同时存在高危的病毒/木马攻击]。整体安全风险主要集中在 [例如：对暴露服务器的攻击以及内部网络的潜在威胁]。")

        document.add_paragraph(f"11.2 **识别出的主要安全风险：**")
        document.add_paragraph("    - [风险1：例如，大规模的外部探测活动，可能为后续攻击做准备]")
        document.add_paragraph("    - [风险2：例如，高危威胁（如xxx）的出现，对关键资产构成威胁]")
        document.add_paragraph("    - [风险3：例如，内部用户可能面临的钓鱼攻击风险]")

        document.add_paragraph(f"11.3 **改进建议：**")
        document.add_paragraph("    - **加强防火墙策略：** 限制对高风险端口 ([例如：3389, 22]) 的非必要访问，特别是来自非信任区域的流量。")
        document.add_paragraph("    - **安全意识培训：** 针对所有员工加强钓鱼邮件、恶意软件预防及安全上网的培训。")
        document.add_paragraph("    - **端点防护：** 确保所有终端设备安装并更新最新版本的防病毒软件/EDR解决方案。")
        document.add_paragraph("    - **日志审计策略：** 审查并优化安全设备的日志审计策略，确保关键事件得到充分记录。")

        document.add_paragraph("---")
        document.add_paragraph("感谢查阅！")

        # 5. 保存文档
        document.save(output_docx_file)
        print(f"安全分析报告已成功生成：'{output_docx_file}'")

    except FileNotFoundError:
        print(f"错误：文件 '{excel_file}' 未找到，请检查路径。")
    except ImportError as e:
        print(f"错误：缺少必要的库: {e}. 请运行 'pip install pandas openpyxl matplotlib seaborn python-docx xlrd'")
    except Exception as e:
        print(f"处理过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

# --- 运行脚本 ---
if __name__ == "__main__":
    generate_security_report(EXCEL_FILE_PATH, REPORT_OUTPUT_PATH)
